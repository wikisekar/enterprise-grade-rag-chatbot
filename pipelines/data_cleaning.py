from pathlib import Path
import hashlib
import shutil

from docx import Document
from pypdf import PdfReader
from pypdf.errors import FileNotDecryptedError


# =====================================================
# Folder Paths
# =====================================================

RAW_FOLDER = Path("data/raw")

DUPLICATE_FOLDER = Path("data/duplicates")
UNSUPPORTED_FOLDER = Path("data/unsupported")
CORRUPTED_FOLDER = Path("data/corrupted")
ENCRYPTED_FOLDER = Path("data/encrypted")
NO_TEXT_FOLDER = Path("data/no_text")

DUPLICATE_FOLDER.mkdir(parents=True, exist_ok=True)
UNSUPPORTED_FOLDER.mkdir(parents=True, exist_ok=True)
CORRUPTED_FOLDER.mkdir(parents=True, exist_ok=True)
ENCRYPTED_FOLDER.mkdir(parents=True, exist_ok=True)
NO_TEXT_FOLDER.mkdir(parents=True, exist_ok=True)


# =====================================================
# Tracked-File Helper
# =====================================================
#
# A file is "tracked" if it's already recorded in
# processed_files.json from a previous successful run.
# Tracked files are trusted -- they already passed every
# cleaning check the first time they were processed, so
# they should not be re-validated (and potentially moved
# out of data/raw/) on every subsequent run. Only brand
# new, untracked files need to run the full check gauntlet.
# =====================================================

def is_tracked(file_path, processed_ids):
    return str(Path(file_path).resolve()) in processed_ids


# =====================================================
# Missing / Empty File Cleaning
# =====================================================
#
# NOTE: This check still runs on ALL files, tracked or not.
# If a tracked file has genuinely vanished from disk or been
# zeroed out, that's a real problem we want to surface, not
# something we should paper over by skipping the check.
# =====================================================

def remove_missing_empty_files(raw_folder=RAW_FOLDER):

    valid_documents = []

    empty_files = 0

    for file in raw_folder.rglob("*"):

        if not file.is_file():
            continue

        if file.stat().st_size == 0:

            print(f"Empty File : {file.name}")

            empty_files += 1

            continue

        valid_documents.append(file)

    print("\n========== Missing / Empty File Report ==========")
    print(f"Valid Files : {len(valid_documents)}")
    print(f"Empty Files : {empty_files}")
    print("=================================================")

    return valid_documents


# =====================================================
# SHA-256 Hash Generator
# =====================================================

def calculate_file_hash(file_path):

    sha256 = hashlib.sha256()

    with open(file_path, "rb") as file:

        while True:

            chunk = file.read(8192)

            if not chunk:
                break

            sha256.update(chunk)

    return sha256.hexdigest()


# =====================================================
# Duplicate Document Cleaning
# =====================================================
#
# If two files share the same SHA-256 hash, always keep
# whichever one is already tracked in processed_files.json
# (even if that means moving a brand-new file into
# DUPLICATE_FOLDER instead of the old one). This prevents a
# previously processed file from silently disappearing from
# disk just because a duplicate showed up later.
# =====================================================

def remove_duplicate_documents(documents, processed_ids=None):

    if processed_ids is None:
        processed_ids = set()

    unique_documents = []

    # hash -> file currently being kept for that hash
    kept_file_by_hash = {}

    duplicate_count = 0

    for file in documents:

        file_hash = calculate_file_hash(file)

        if file_hash not in kept_file_by_hash:

            kept_file_by_hash[file_hash] = file

            unique_documents.append(file)

            continue

        # Hash collision: decide which of the two to keep.
        existing_file = kept_file_by_hash[file_hash]

        current_tracked = is_tracked(file, processed_ids)
        existing_tracked = is_tracked(existing_file, processed_ids)

        if current_tracked and not existing_tracked:

            # The file we already kept is NOT tracked, but this
            # new collision IS tracked. Swap: keep the tracked one,
            # move the previously-kept untracked one instead.

            print(
                f"Duplicate Document : {existing_file.name} "
                f"(swapped out to keep already-tracked file: {file.name})"
            )

            shutil.move(
                str(existing_file),
                DUPLICATE_FOLDER / existing_file.name
            )

            unique_documents.remove(existing_file)
            unique_documents.append(file)

            kept_file_by_hash[file_hash] = file

        else:

            # Default behavior: keep whichever file we already kept
            # (this also covers the case where neither, or both,
            # are tracked -- in which case original behavior stands).

            print(f"Duplicate Document : {file.name}")

            shutil.move(
                str(file),
                DUPLICATE_FOLDER / file.name
            )

        duplicate_count += 1

    print("\n========== Duplicate Document Report ==========")
    print(f"Unique Documents    : {len(unique_documents)}")
    print(f"Duplicate Documents : {duplicate_count}")
    print("================================================")

    return unique_documents


# =====================================================
# Unsupported File Type Cleaning
# =====================================================

def remove_unsupported_file_types(documents):

    allowed_extensions = {

        ".pdf",
        ".docx",
        ".txt"

    }

    supported_documents = []

    unsupported_count = 0

    for file in documents:

        extension = file.suffix.lower()

        if extension in allowed_extensions:

            supported_documents.append(file)

        else:

            print(f"Unsupported File : {file.name}")

            shutil.move(
                str(file),
                UNSUPPORTED_FOLDER / file.name
            )

            unsupported_count += 1

    print("\n========== Unsupported File Type Report ==========")
    print(f"Supported Documents   : {len(supported_documents)}")
    print(f"Unsupported Documents : {unsupported_count}")
    print("===================================================")

    return supported_documents


# =====================================================
# Corrupted Document Cleaning
# =====================================================

def remove_corrupted_documents(documents):

    valid_documents = []

    corrupted_count = 0

    for file in documents:

        extension = file.suffix.lower()

        try:

            if extension == ".pdf":

                PdfReader(file)

            elif extension == ".docx":

                Document(file)

            elif extension == ".txt":

                with open(
                    file,
                    "r",
                    encoding="utf-8",
                    errors="ignore"
                ) as f:

                    f.read()

            valid_documents.append(file)

        except Exception:

            print(f"Corrupted Document : {file.name}")

            shutil.move(
                str(file),
                CORRUPTED_FOLDER / file.name
            )

            corrupted_count += 1

    print("\n========== Corrupted Document Report ==========")
    print(f"Valid Documents     : {len(valid_documents)}")
    print(f"Corrupted Documents : {corrupted_count}")
    print("================================================")

    return valid_documents


# =====================================================
# Encrypted PDF Cleaning
# =====================================================

def remove_encrypted_documents(documents):

    valid_documents = []

    encrypted_count = 0

    for file in documents:

        if file.suffix.lower() != ".pdf":

            valid_documents.append(file)

            continue

        try:

            reader = PdfReader(file)

            if reader.is_encrypted:

                print(f"Encrypted PDF : {file.name}")

                shutil.move(
                    str(file),
                    ENCRYPTED_FOLDER / file.name
                )

                encrypted_count += 1

            else:

                valid_documents.append(file)

        except Exception:

            valid_documents.append(file)

    print("\n========== Encrypted PDF Report ==========")
    print(f"Readable PDFs  : {len(valid_documents)}")
    print(f"Encrypted PDFs : {encrypted_count}")
    print("==========================================")

    return valid_documents


# =====================================================
# Documents Without Extractable Text
# =====================================================

def remove_documents_without_text(documents):

    valid_documents = []

    no_text_count = 0

    for file in documents:

        extracted_text = ""

        try:

            if file.suffix.lower() == ".pdf":

                reader = PdfReader(file)

                for page in reader.pages:

                    text = page.extract_text()

                    if text:

                        extracted_text += text

            elif file.suffix.lower() == ".docx":

                doc = Document(file)

                for paragraph in doc.paragraphs:

                    extracted_text += paragraph.text

            elif file.suffix.lower() == ".txt":

                with open(
                    file,
                    "r",
                    encoding="utf-8",
                    errors="ignore"
                ) as f:

                    extracted_text = f.read()

            if extracted_text.strip():

                valid_documents.append(file)

            else:

                print(f"No Text : {file.name}")

                shutil.move(
                    str(file),
                    NO_TEXT_FOLDER / file.name
                )

                no_text_count += 1

        except Exception:

            valid_documents.append(file)

    print("\n========== No Extractable Text Report ==========")
    print(f"Valid Documents : {len(valid_documents)}")
    print(f"No Text Files   : {no_text_count}")
    print("================================================")

    return valid_documents


# =====================================================
# Data Cleaning Pipeline
# =====================================================
#
# KEY CHANGE: after duplicate handling, we split the
# remaining documents into two groups:
#
#   - tracked_documents   : already in processed_files.json
#   - untracked_documents : brand new, never seen before
#
# Only untracked_documents go through the unsupported /
# corrupted / encrypted / no-text checks. Tracked documents
# already passed all of that the first time they were
# processed, so they're trusted and pass straight through.
# This closes off the entire class of bug where a cleaning
# step moves an already-tracked file out of data/raw/ and
# silently shrinks the valid-file count below what's in the
# tracker.
# =====================================================

def run_data_cleaning_pipeline(processed_ids=None):

    if processed_ids is None:
        processed_ids = set()

    documents = remove_missing_empty_files()

    documents = remove_duplicate_documents(documents, processed_ids)

    tracked_documents = [
        file for file in documents
        if is_tracked(file, processed_ids)
    ]

    untracked_documents = [
        file for file in documents
        if not is_tracked(file, processed_ids)
    ]

    print(
        f"\n(Skipping unsupported/corrupted/encrypted/no-text checks "
        f"for {len(tracked_documents)} already-tracked file(s); "
        f"running full checks on {len(untracked_documents)} new file(s).)"
    )

    untracked_documents = remove_unsupported_file_types(untracked_documents)

    untracked_documents = remove_corrupted_documents(untracked_documents)

    untracked_documents = remove_encrypted_documents(untracked_documents)

    untracked_documents = remove_documents_without_text(untracked_documents)

    documents = tracked_documents + untracked_documents

    return documents


# =====================================================
# Standalone Test Runner
# =====================================================

if __name__ == "__main__":
    run_data_cleaning_pipeline()